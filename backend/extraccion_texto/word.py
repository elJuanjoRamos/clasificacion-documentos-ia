import os
import win32com.client

from docx import Document

from ..ocr.ocr import extract_ocr_image
from .utils import (
    clean_temp_dir,
    clean_temp_files,
    clean_text,
    create_block,
    build_document_output,
    normalize_score,
    score_to_confidence,
    score_to_quality
)

OUTPUT_DIR = "temp_files"


def main_word(ext, path):
    blocks = []
    images = []
    error = None
    supported = False

    try:
        working_path, extraction_method = prepare_word_file(ext, path)

        #ABRE EL DOCUMENTO
        doc = Document(working_path)
        # EXTRAE LOS BLOQUES DE TEXTO
        blocks = get_word_text(doc)
        # EXTRAE IMAGENES
        images = get_images_from_word(doc, OUTPUT_DIR)

        ocr_blocks = apply_ocr_to_word_images(
            image_paths=images,
            start_order=len(blocks) + 1
        )

        blocks.extend(ocr_blocks)

        supported = True

    except Exception as e:
        extraction_method = None
        error = str(e)
    finally:
        clean_temp_files(images)
        clean_temp_dir(OUTPUT_DIR)

    return build_document_output(
        file_type="word",
        extension=ext,
        supported=supported,
        extraction_method=extraction_method,
        blocks=blocks,
        error=error,
        file_path=path,
        metadata={
            "images": images
        },
        type_metadata_key="word_metadata"
    )


def prepare_word_file(ext: str, path: str) -> tuple[str, str]:
    if ext == ".doc":
        os.makedirs(OUTPUT_DIR, exist_ok=True)

        input_path = os.path.abspath(str(path))

        filename = os.path.splitext(os.path.basename(input_path))[0]
        output_path = os.path.abspath(
            os.path.join(OUTPUT_DIR, f"{filename}.docx")
        )

        word = win32com.client.Dispatch("Word.Application")
        word.Visible = False

        doc = None

        try:
            doc = word.Documents.Open(input_path)
            doc.SaveAs(output_path, FileFormat=16)
            return output_path, "doc_converted_to_docx"

        finally:
            if doc is not None:
                doc.Close(False)

            word.Quit()

    return str(path), "docx_structured"



def get_word_text(doc: Document) -> list[dict]:
    blocks = []

    blocks.extend(extract_paragraph_blocks(doc))

    blocks = merge_related_paragraphs(blocks)

    blocks.extend(
        extract_table_blocks(
            doc,
            start_order=len(blocks) + 1
        )
    )

    return blocks


def extract_paragraph_blocks(doc: Document) -> list[dict]:
    blocks = []

    for paragraph in doc.paragraphs:
        text = clean_text(paragraph.text)

        if not text:
            continue

        style_name = paragraph.style.name if paragraph.style else ""

        block_type = detect_paragraph_type(style_name)

        if block_type == "list_item":
            text = f"- {text}"

        blocks.append(create_block(
            order=len(blocks) + 1,
            block_type=block_type,
            text=text,
            source="word_paragraph",
            metadata={
                "style": style_name
            }
        ))

    return blocks


def detect_paragraph_type(style_name: str) -> str:
    style = style_name.lower()

    if "heading" in style or "title" in style:
        return "title"

    if "list" in style:
        return "list_item"

    return "paragraph"


def extract_table_blocks(doc: Document, start_order: int) -> list[dict]:
    blocks = []

    for table_index, table in enumerate(doc.tables, start=1):
        rows = []

        for row in table.rows:
            cells = [
                clean_text(cell.text)
                for cell in row.cells
            ]

            if any(cells):
                rows.append(cells)

        if not rows:
            continue

        column_count = max(len(row) for row in rows)

        rows = [
            row + [""] * (column_count - len(row))
            for row in rows
        ]

        header = rows[0]
        body = rows[1:]

        markdown_lines = [
            "| " + " | ".join(header) + " |",
            "| " + " | ".join(["---"] * column_count) + " |"
        ]

        for row in body:
            markdown_lines.append(
                "| " + " | ".join(row) + " |"
            )

        blocks.append(create_block(
            order=start_order + len(blocks),
            block_type="table",
            text="\n".join(markdown_lines),
            source="word_table",
            metadata={
                "table_index": table_index,
                "row_count": len(rows),
                "column_count": column_count
            }
        ))

    return blocks




def get_images_from_word(doc: Document, output_dir: str) -> list[str]:
    image_paths = []

    os.makedirs(output_dir, exist_ok=True)

    for index, rel in enumerate(doc.part.rels.values(), start=1):
        if "image" not in rel.target_ref:
            continue

        _, ext = os.path.splitext(rel.target_ref)

        ext = ext.lower() if ext else ".png"

        image_path = os.path.join(
            output_dir,
            f"word_img_{index}{ext}"
        )

        with open(image_path, "wb") as file:
            file.write(rel.target_part.blob)

        image_paths.append(image_path)

    return image_paths


def apply_ocr_to_word_images(
    image_paths: list[str],
    start_order: int
) -> list[dict]:

    blocks = []
    text = ""
    for image_index, image_path in enumerate(image_paths, start=1):
        try:
            result = extract_ocr_image(image_path)

            for ocr_block in result.get("blocks", []):
                text = clean_text(ocr_block.get("text", ""))

                if not text:
                    continue

                original_metadata = ocr_block.get("metadata", {}) or {}
                raw_ocr_score = original_metadata.get("score", 0.0)
                normalized_score = normalize_score(raw_ocr_score, max_value=2.0)
                ocr_confidence = original_metadata.get("confidence", 0.0)

                blocks.append(create_block(
                    order=start_order + len(blocks),
                    block_type="image_ocr",
                    text=text,
                    source="word_image",
                    metadata={
                        "image_index": image_index,
                        "file_path": image_path,
                        "ocr": {
                            "raw_score": raw_ocr_score,
                            "raw_quality": original_metadata.get("quality"),
                            "ocr_confidence": ocr_confidence,
                            "global_result": result.get("ocr", {})
                        },
                        "preprocessing": result.get("preprocessing", {}),
                        "bbox": original_metadata.get("bbox"),
                    },
                    precomputed_quality={
                        "score": normalized_score,
                        "confidence": score_to_confidence(normalized_score),
                        "quality": score_to_quality(normalized_score),
                        "accepted": original_metadata.get("accepted", True),
                        "rejection_reason": original_metadata.get("rejection_reason"),
                        "metrics": {
                            "text_length": original_metadata.get("text_length", len(text)),
                            "ocr_confidence": ocr_confidence,
                            "raw_ocr_score": raw_ocr_score,
                            "raw_ocr_quality": original_metadata.get("quality")
                        }
                    }
                ))

        except Exception as e:
            blocks.append(create_block(
                order=start_order + len(blocks),
                block_type="image_ocr",
                text="",
                source="word_image",
                metadata={
                    "image_index": image_index,
                    "file_path": image_path,
                    "ocr": {
                        "accepted": False,
                        "rejection_reason": "ocr_error"
                    },
                    "error": str(e)
                }
            ))

    return blocks

def merge_related_paragraphs(blocks: list[dict]) -> list[dict]:
    merged = []

    buffer = None

    for block in blocks:
        if block.get("type") not in ["paragraph", "title", "list_item"]:
            if buffer:
                merged.append(buffer)
                buffer = None

            merged.append(block)
            continue

        if buffer is None:
            buffer = block.copy()
            continue

        if should_merge_paragraphs(buffer, block):
            buffer["text"] = (
                f'{buffer["text"].rstrip()} '
                f'{block["text"].lstrip()}'
            )
        else:
            merged.append(buffer)
            buffer = block.copy()

    if buffer:
        merged.append(buffer)

    recalculated = []

    for index, block in enumerate(merged, start=1):
        if block.get("type") in ["paragraph", "title", "list_item"]:
            block = create_block(
                order=index,
                block_type=block.get("type"),
                text=block.get("text", ""),
                source=block.get("source", ""),
                metadata=block.get("metadata", {})
            )
        else:
            block["order"] = index

        recalculated.append(block)

    return recalculated


def should_merge_paragraphs(
    current: dict,
    next_block: dict
) -> bool:

    current_text = current.get("text", "").strip()
    next_text = next_block.get("text", "").strip()

    if current.get("type") == "list_item":
        if current_text.endswith((".", ":", ";", "?", "!")):
            return False

        if next_text and next_text[0].islower():
            return True

        return False

    if not current_text or not next_text:
        return False

    if current.get("type") != next_block.get("type"):
        return False

    current_style = current.get(
        "metadata",
        {}
    ).get("style", "")

    next_style = next_block.get(
        "metadata",
        {}
    ).get("style", "")

    if current_style != next_style:
        return False

    strong_endings = (
        ".", ":", ";", "?", "!"
    )

    continuation_words = {
        "y", "e", "o", "u",
        "de", "del", "la", "el",
        "los", "las", "en",
        "para", "con", "por", "a", 
        "and", "or", "not"
    }

    last_word = current_text.split()[-1].lower()

    if last_word in continuation_words:
        return True

    return (
        len(current_text) <= 90
        and len(next_text) <= 90
        and not current_text.endswith(strong_endings)
    )



